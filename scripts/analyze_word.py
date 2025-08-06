#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析Word文档中的表格和项目提取情况
"""

from docx import Document
import re

def analyze_word():
    """分析Word文档中的表格和项目提取情况"""
    docx_path = "docs/儿童生长发育量表.docx"
    doc = Document(docx_path)
    
    print(f"Word文档包含 {len(doc.tables)} 个表格")
    
    total_projects = 0
    all_project_ids = []
    
    for i, table in enumerate(doc.tables):
        print(f"\n表格 {i+1}: {len(table.rows)} 行 x {len(table.columns)} 列")
        
        # 检查表格内容
        table_projects = 0
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if '□' in cell_text:
                    matches = re.findall(r'□(\d+)', cell_text)
                    table_projects += len(matches)
                    all_project_ids.extend(matches)
        
        print(f"  表格 {i+1} 中的项目数: {table_projects}")
        total_projects += table_projects
    
    print(f"\n总计项目数: {total_projects}")
    print(f"唯一项目ID数: {len(set(all_project_ids))}")
    
    if all_project_ids:
        project_ids = [int(id) for id in all_project_ids if id.isdigit()]
        print(f"项目ID范围: {min(project_ids)} - {max(project_ids)}")
        print(f"缺失的项目ID: {set(range(1, max(project_ids) + 1)) - set(project_ids)}")

if __name__ == "__main__":
    analyze_word() 