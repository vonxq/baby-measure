#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档表格数据提取脚本
从儿童生长发育量表.docx中提取表格数据
"""

import json
import re
from docx import Document
import pandas as pd
from pathlib import Path

class ScaleDataExtractor:
    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.items = []
        self.methods = []
        self.extracted_ids = set()  # 用于去重
        
    def extract_table_data(self):
        """提取所有表格数据"""
        print(f"正在处理文档: {self.docx_path}")
        print(f"文档包含 {len(self.doc.tables)} 个表格")
        
        for i, table in enumerate(self.doc.tables):
            print(f"\n处理表格 {i+1}:")
            print(f"表格大小: {len(table.rows)} 行 x {len(table.columns)} 列")
            
            # 提取表格数据
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
                print(f"行数据: {row_data}")
            
            # 检查表格是否包含项目数据
            has_projects = self._check_has_projects(table_data)
            print(f"包含项目数据: {has_projects}")
            
            # 尝试识别表格类型
            if self._is_scale_table(table_data) or has_projects:
                print("识别为量表项目表格")
                self._extract_scale_items(table_data)
            elif self._is_method_table(table_data):
                print("识别为操作方法表格")
                self._extract_methods(table_data)
            else:
                print("未知表格类型")
    
    def _check_has_projects(self, table_data):
        """检查表格是否包含项目数据"""
        for row in table_data:
            for cell in row:
                if '□' in cell:
                    return True
        return False
    
    def _is_scale_table(self, table_data):
        """判断是否为量表项目表格"""
        if not table_data:
            return False
        
        # 检查是否包含月龄信息
        first_row = ' '.join(table_data[0])
        return any(month in first_row for month in ['月龄', '1月', '2月', '3月', '4月', '5月', '6月'])
    
    def _is_method_table(self, table_data):
        """判断是否为操作方法表格"""
        if not table_data:
            return False
        
        # 检查是否包含操作方法相关字段
        first_row = ' '.join(table_data[0])
        return any(keyword in first_row for keyword in ['操作方法', '测查通过要求'])
    
    def _extract_scale_items(self, table_data):
        """提取量表项目数据"""
        print("开始提取量表项目...")
        
        # 解析表头，获取月龄信息
        headers = table_data[0] if table_data else []
        print(f"表头: {headers}")
        
        # 首先从表头中提取项目（如果有的话）
        self._extract_items_from_row(headers, headers, 0, is_header=True)
        
        # 解析项目数据
        for row_idx, row in enumerate(table_data[1:], 1):
            if not row or not any(cell.strip() for cell in row):
                continue
                
            # 尝试解析项目
            self._extract_items_from_row(row, headers, row_idx)
    
    def _extract_items_from_row(self, row, headers, row_idx, is_header=False):
        """从行中提取项目"""
        for cell_idx, cell in enumerate(row):
            if not cell.strip():
                continue
                
            # 查找所有项目编号和名称（支持一个单元格多个项目）
            item_matches = re.findall(r'□(\d+)\s*([^□]+)', cell)
            if item_matches:
                # 计算该单元格中项目的分值平分
                total_score = self._calculate_score_for_column(headers, cell_idx)
                score_per_item = total_score / len(item_matches) if item_matches else 1.0
                
                for item_id, item_name in item_matches:
                    # 去重检查
                    if item_id in self.extracted_ids:
                        print(f"跳过重复项目: {item_id} - {item_name.strip()}")
                        continue
                    
                    item_name = item_name.strip()
                    
                    # 确定月龄组
                    age_group = self._determine_age_group(headers, cell_idx)
                    
                    # 确定发育领域
                    category = self._determine_category_from_row(row, row_idx)
                    
                    # 检查特殊标记
                    is_important = '*' in item_name
                    can_ask_parent = 'R' in item_name
                    
                    # 清理项目名称
                    clean_name = item_name.replace('*', '').replace('R', '').strip()
                    
                    item_data = {
                        "id": item_id,
                        "category": category,
                        "age_group": age_group,
                        "item_name": clean_name,
                        "description": "",  # 需要从操作方法表获取
                        "pass_criteria": "",  # 需要从操作方法表获取
                        "score": score_per_item,
                        "is_important": is_important,
                        "can_ask_parent": can_ask_parent
                    }
                    
                    self.items.append(item_data)
                    self.extracted_ids.add(item_id)  # 添加到已提取集合
                    print(f"提取项目: {item_id} - {clean_name} (分值: {score_per_item})")
    
    def _extract_methods(self, table_data):
        """提取操作方法数据"""
        print("开始提取操作方法...")
        
        for row in table_data[1:]:  # 跳过表头
            if len(row) < 3:
                continue
                
            # 尝试解析项目编号
            item_id_match = re.search(r'(\d+)', row[0])
            if item_id_match:
                item_id = item_id_match.group(1)
                
                method_data = {
                    "item_id": item_id,
                    "operation_method": row[1] if len(row) > 1 else "",
                    "pass_requirement": row[2] if len(row) > 2 else ""
                }
                
                self.methods.append(method_data)
                print(f"提取方法: {item_id}")
    
    def _determine_age_group(self, headers, column_idx):
        """确定月龄组"""
        if column_idx < len(headers):
            header = headers[column_idx]
            # 提取月龄信息
            month_match = re.search(r'(\d+)月龄', header)
            if month_match:
                return f"{month_match.group(1)}个月"
        return "未知月龄"
    
    def _determine_category_from_row(self, row, row_idx):
        """从行内容确定发育领域"""
        row_text = ' '.join(row)
        if '大运动' in row_text or '大 运 动' in row_text:
            return "大运动"
        elif '精细动作' in row_text or '精细动作' in row_text:
            return "精细动作"
        elif '适应能力' in row_text:
            return "适应能力"
        elif '语言' in row_text or '语    言' in row_text:
            return "语言"
        elif '社会行为' in row_text:
            return "社会行为"
        else:
            return "未知领域"
    
    def _calculate_score_for_column(self, headers, column_idx):
        """计算指定列的分值"""
        if column_idx < len(headers):
            header = headers[column_idx]
            # 提取月龄信息
            month_match = re.search(r'(\d+)月龄', header)
            if month_match:
                month = int(month_match.group(1))
                if 1 <= month <= 12:
                    return 1.0
                elif 15 <= month <= 36:
                    return 3.0
                elif 42 <= month <= 84:
                    return 6.0
        return 1.0
    
    def _calculate_score(self, age_group):
        """计算分值"""
        if '个月' in age_group:
            month = int(age_group.replace('个月', ''))
            if 1 <= month <= 12:
                return 1.0
            elif 15 <= month <= 36:
                return 3.0
            elif 42 <= month <= 84:
                return 6.0
        return 1.0
    
    def merge_data(self):
        """合并项目和方法数据"""
        print("\n开始合并数据...")
        
        # 创建项目ID到方法的映射
        method_map = {method['item_id']: method for method in self.methods}
        
        # 更新项目数据
        for item in self.items:
            if item['id'] in method_map:
                method = method_map[item['id']]
                item['description'] = method['operation_method']
                item['pass_criteria'] = method['pass_requirement']
        
        print(f"合并完成，共 {len(self.items)} 个项目")
    
    def save_data(self, output_dir):
        """保存提取的数据"""
        output_path = Path(output_dir)
        output_path.mkdir(exist_ok=True)
        
        # 保存项目数据
        items_file = output_path / "assessment_items.json"
        with open(items_file, 'w', encoding='utf-8') as f:
            json.dump({"items": self.items}, f, ensure_ascii=False, indent=2)
        print(f"项目数据已保存到: {items_file}")
        
        # 保存方法数据
        methods_file = output_path / "assessment_methods.json"
        with open(methods_file, 'w', encoding='utf-8') as f:
            json.dump({"methods": self.methods}, f, ensure_ascii=False, indent=2)
        print(f"方法数据已保存到: {methods_file}")
        
        # 生成统计信息
        print(f"\n提取统计:")
        print(f"- 项目总数: {len(self.items)}")
        print(f"- 方法总数: {len(self.methods)}")
        
        # 按发育领域统计
        categories = {}
        for item in self.items:
            cat = item['category']
            categories[cat] = categories.get(cat, 0) + 1
        
        print(f"- 按发育领域分布:")
        for cat, count in categories.items():
            print(f"  {cat}: {count} 项")
        
        # 按月龄组统计
        age_groups = {}
        for item in self.items:
            age = item['age_group']
            age_groups[age] = age_groups.get(age, 0) + 1
        
        print(f"- 按月龄组分布:")
        for age, count in age_groups.items():
            print(f"  {age}: {count} 项")

def main():
    """主函数"""
    docx_path = "docs/儿童生长发育量表.docx"
    
    if not Path(docx_path).exists():
        print(f"错误: 文件不存在 {docx_path}")
        return
    
    try:
        # 创建提取器
        extractor = ScaleDataExtractor(docx_path)
        
        # 提取数据
        extractor.extract_table_data()
        
        # 合并数据
        extractor.merge_data()
        
        # 保存数据
        extractor.save_data("data")
        
        print("\n✅ 数据提取完成!")
        
    except Exception as e:
        print(f"❌ 提取失败: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main() 